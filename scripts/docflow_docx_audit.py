import json
import re
import sys
import zipfile
from collections import Counter
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}


def qn(name: str) -> str:
    prefix, local = name.split(":")
    return f"{{{NS[prefix]}}}{local}"


def para_text(p) -> str:
    return "".join(t.text or "" for t in p._p.findall(".//w:t", NS)).strip()


def detect_heading_level(paragraph) -> int | None:
    style_name = (paragraph.style.name or "") if paragraph.style else ""
    m = re.search(r"(?:Heading|标题)\s*(\d+)", style_name, re.I)
    if m:
        level = int(m.group(1))
        if 1 <= level <= 9:
            return level

    ppr = paragraph._p.find("w:pPr", NS)
    if ppr is not None:
        outline = ppr.find("w:outlineLvl", NS)
        if outline is not None:
            val = outline.get(qn("w:val"))
            if val and val.isdigit():
                return int(val) + 1

    text = para_text(paragraph)
    if not text or len(text) > 90:
        return None
    patterns = [
        (1, r"^第[一二三四五六七八九十百千万\d]+[章节篇部]\b"),
        (1, r"^[一二三四五六七八九十]+[、．.]\s*[^。；;]{2,}$"),
        (1, r"^\d+[、．.]\s*[^。；;]{2,}$"),
        (2, r"^（[一二三四五六七八九十]+）\s*[^。；;]{2,}$"),
        (2, r"^\d+\.\d+\s+[^。；;]{2,}$"),
        (3, r"^（\d+）\s*[^。；;]{2,}$"),
        (3, r"^\d+\.\d+\.\d+\s+[^。；;]{2,}$"),
    ]
    for level, pattern in patterns:
        if re.match(pattern, text):
            return level
    return None


def extract_docx_xml_counts(path: Path) -> dict:
    with zipfile.ZipFile(path) as z:
        document_xml = z.read("word/document.xml")
        root = ET.fromstring(document_xml)
        rels = []
        try:
            rels_xml = z.read("word/_rels/document.xml.rels")
            rels_root = ET.fromstring(rels_xml)
            rels = rels_root.findall("{http://schemas.openxmlformats.org/package/2006/relationships}Relationship")
        except KeyError:
            pass

    image_rels = [
        r for r in rels
        if (r.get("Type") or "").endswith("/image")
    ]
    drawings = root.findall(".//w:drawing", NS)
    picts = root.findall(".//w:pict", NS)
    tables = root.findall(".//w:tbl", NS)
    return {
        "image_relationships": len(image_rels),
        "drawing_elements": len(drawings),
        "legacy_picture_elements": len(picts),
        "xml_tables": len(tables),
    }


def audit_docx(path: Path) -> dict:
    doc = Document(str(path))
    paragraphs = [p for p in doc.paragraphs if para_text(p)]
    headings = []
    figure_captions = []
    table_captions = []
    other_captions = []
    for idx, p in enumerate(doc.paragraphs):
        text = para_text(p)
        if not text:
            continue
        level = detect_heading_level(p)
        if level:
            headings.append({
                "index": idx,
                "level": level,
                "style": p.style.name if p.style else "",
                "text": text[:160],
            })
        normalized = re.sub(r"\s+", "", text)
        if re.match(r"^图\s*\d+([-.]\d+)?", normalized):
            figure_captions.append(text[:160])
        elif re.match(r"^表\s*\d+([-.]\d+)?", normalized):
            table_captions.append(text[:160])
        elif re.match(r"^(图|表)\s*[一二三四五六七八九十]+", normalized):
            other_captions.append(text[:160])

    heading_counter = Counter(h["level"] for h in headings)
    xml_counts = extract_docx_xml_counts(path)
    char_count = sum(len(para_text(p)) for p in doc.paragraphs)
    result = {
        "file": str(path),
        "size_bytes": path.stat().st_size,
        "paragraphs": len(paragraphs),
        "char_count": char_count,
        "tables_python_docx": len(doc.tables),
        **xml_counts,
        "headings": len(headings),
        "headings_by_level": dict(sorted(heading_counter.items())),
        "figure_captions": len(figure_captions),
        "table_captions": len(table_captions),
        "other_captions": len(other_captions),
        "heading_samples": headings[:40],
        "figure_caption_samples": figure_captions[:20],
        "table_caption_samples": table_captions[:20],
    }
    return result


def main() -> None:
    if len(sys.argv) < 2:
        raise SystemExit("usage: docflow_docx_audit.py input.docx [out.json]")
    path = Path(sys.argv[1])
    result = audit_docx(path)
    text = json.dumps(result, ensure_ascii=False, indent=2)
    if len(sys.argv) >= 3:
        Path(sys.argv[2]).write_text(text, encoding="utf-8")
    print(text)


if __name__ == "__main__":
    main()
